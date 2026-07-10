from pathlib import Path
from typing import List, Dict, Any
from parsers.base_parser import BaseParser
from utils.logger import get_logger

logger = get_logger(__name__)


class DOCXParser(BaseParser):
    """Parse DOCX files using python-docx."""

    def parse(self, file_path: str) -> List[Dict[str, Any]]:
        try:
            from docx import Document
        except ImportError:
            logger.error("python-docx not installed. Run: pip install python-docx")
            return []

        filename = Path(file_path).name
        results = []

        try:
            doc = Document(file_path)
            full_text = []

            for para in doc.paragraphs:
                if para.text.strip():
                    full_text.append(para.text.strip())

            # Also extract tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(
                        cell.text.strip() for cell in row.cells if cell.text.strip()
                    )
                    if row_text:
                        full_text.append(row_text)

            combined = "\n".join(full_text)
            if combined.strip():
                results.append(
                    {
                        "text": combined.strip(),
                        "metadata": {
                            "filename": filename,
                            "page": 1,
                            "sheet": None,
                        },
                    }
                )
            else:
                logger.warning(f"DOCX has no extractable text: {filename}")

        except Exception as e:
            logger.error(f"Failed to parse DOCX {filename}: {e}")
            return []

        logger.info(f"DOCX parsed: {filename} ({len(results)} blocks extracted)")
        return results
